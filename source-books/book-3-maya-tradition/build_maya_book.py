#!/usr/bin/env python3
"""Build mobile-readable reading editions from the source-backed Maya manuscript."""

from __future__ import annotations

import html
import json
import re
import shutil
from collections import defaultdict
from html.parser import HTMLParser
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError:
    Document = None


HERE = Path(__file__).resolve().parent
MANUSCRIPT = HERE / "manuscript" / "MAYA_TRADITION.md"
UNIFIED_MANUSCRIPT = HERE / "manuscript" / "MAYA_TRADITION_UNIFIED.md"
RAW_MESSAGES = HERE / "raw" / "messages.html"
RAW_PHOTOS = HERE / "raw" / "photos"
SUPPLEMENTAL_INDEX = HERE / "raw" / "templetherapy" / "TEMPLETHERAPY_MAYA_AZTEC_INDEX.jsonl"
SUPPLEMENTAL_MEDIA = HERE / "media" / "templetherapy"
OUT = HERE / "outputs"
HTML_OUT = OUT / "Maya_Tradition_Methodology.html"
DOCX_OUT = OUT / "Maya_Tradition_Methodology.docx"
DEDUPLICATION_REPORT = HERE / "DEDUPLICATION_REPORT.md"

WARM = "6B2F1A"
GOLD = "B7791F"
CREAM = "FBF4E9"
INK = "2F241D"
MUTED = "6D5A4D"
FRONT_HEADINGS = {"Editorial note", "Описание традиции", "Содержание", "Авторская рамка, практики и программы"}
CHAPTERS = (
    "I. Описание традиции",
    "II. Боги и божественные силы",
    "III. Календарь, время и космология",
    "IV. Места, предметы и материальная культура",
    "V. Мифология, Шибальба, инициация и ритуал",
    "VI. Авторские, архетипические и терапевтические модели",
    "VII. Сравнения мезоамериканских традиций",
)

# These are editorial placement decisions, not claims about a source's authority.
SUPPLEMENTAL_CHAPTERS = {
    2062: CHAPTERS[0], 2065: CHAPTERS[2], 2100: CHAPTERS[5], 2113: CHAPTERS[3],
    2198: CHAPTERS[5], 2204: CHAPTERS[5], 2210: CHAPTERS[5], 2212: CHAPTERS[5],
    2217: CHAPTERS[1], 2239: CHAPTERS[4], 2240: CHAPTERS[5], 2241: CHAPTERS[4],
    2242: CHAPTERS[4], 2245: CHAPTERS[4], 2246: CHAPTERS[4], 2251: CHAPTERS[4],
    2253: CHAPTERS[5], 2255: CHAPTERS[5], 2257: CHAPTERS[2], 2258: CHAPTERS[2],
    2262: CHAPTERS[0], 2264: CHAPTERS[5], 2268: CHAPTERS[1], 2269: CHAPTERS[5],
    2273: CHAPTERS[5], 2323: CHAPTERS[3], 2346: CHAPTERS[1], 2352: CHAPTERS[6],
    2446: CHAPTERS[5],
}

# Supplemental copies are linked from the retained Mayaismagic article.  The
# three near copies were reviewed separately: 2246 adds a tail, 2268 changes a
# short closing sentence, and 2346 reorders the same material.
SUPPLEMENTAL_CANONICAL = {
    2065: 89, 2113: 98, 2204: 147, 2217: 214, 2239: 222, 2240: 223,
    2241: 224, 2242: 225, 2245: 226, 2246: 227, 2251: 153, 2255: 154,
    2257: 155, 2258: 156, 2268: 161, 2323: 240, 2346: 245,
}


def parse_supplemental_articles() -> list[dict[str, object]]:
    """Read substantive public TempleTherapy entries with namespace-safe IDs."""
    articles: list[dict[str, object]] = []
    for line_number, line in enumerate(SUPPLEMENTAL_INDEX.read_text(encoding="utf-8").splitlines(), 1):
        entry = json.loads(line)
        raw_text = html.unescape(str(entry["raw_text"]))
        if not raw_text.strip():
            raise ValueError(f"TempleTherapy post at line {line_number} has no substantive text")
        articles.append({
            "chapter": SUPPLEMENTAL_CHAPTERS[int(entry["post_id"])],
            "channel": "TempleTherapy",
            "title": f"TempleTherapy · пост {entry['post_id']}",
            "post_id": entry["post_id"],
            "article_id": f"templetherapy-{entry['post_id']}",
            "url": entry["url"],
            "date": entry["date"],
            "text": raw_text,
            "media_references": entry["media_references"],
            "media_root": "media/templetherapy",
        })
    return articles


def parse_media() -> dict[int, list[str]]:
    class MediaParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.media: dict[int, list[str]] = {}
            self.div_stack: list[int | None] = []

        def handle_starttag(self, tag, attrs):
            attributes = dict(attrs)
            if tag == "div":
                classes = set(attributes.get("class", "").split())
                matched = re.fullmatch(r"message(\d+)", attributes.get("id", ""))
                post_id = int(matched.group(1)) if {"message", "default"} <= classes and matched else None
                self.div_stack.append(post_id)
                if post_id is not None:
                    self.media[post_id] = []
                return
            post_id = next((item for item in reversed(self.div_stack) if item is not None), None)
            if post_id is None or tag not in {"a", "img"}:
                return
            ref = attributes.get("href") or attributes.get("src")
            if ref and ref.startswith("photos/") and ref not in self.media[post_id]:
                self.media[post_id].append(ref)

        def handle_endtag(self, tag):
            if tag == "div" and self.div_stack:
                self.div_stack.pop()

    parser = MediaParser()
    parser.feed(RAW_MESSAGES.read_text(encoding="utf-8"))
    return parser.media


def parse_source(body: str) -> tuple[int, str, str, str] | None:
    source = re.search(r"^\s*\*Источник: пост \[(\d+)\]\((https://t\.me/[^)]+)\); (.+)\.\*\s*\n", body)
    if not source:
        return None
    return int(source.group(1)), source.group(2), source.group(3), body[source.end():].strip()


def parse_front_description(text: str) -> dict[str, object]:
    block = re.search(r"(?ms)^# Описание традиции\n\n(.*?)(?=^# Содержание)", text)
    if not block:
        raise ValueError("Missing source-derived front description")
    heading = re.search(r"(?m)^## (.+)$", block.group(1))
    if not heading:
        raise ValueError("Missing front description source heading")
    source = parse_source(block.group(1)[heading.end():])
    if not source:
        raise ValueError("Missing front description citation")
    post_id, url, date, body = source
    return {"title": heading.group(1), "post_id": post_id, "url": url, "date": date, "text": body}


def parse_articles() -> list[dict[str, object]]:
    text = MANUSCRIPT.read_text(encoding="utf-8")
    front_end = text.index("# Содержание")
    current_chapter = ""
    articles: list[dict[str, object]] = []
    headings = list(re.finditer(r"(?m)^# (.+)$|^## (.+)$", text))
    for index, match in enumerate(headings):
        level, heading = (1, match.group(1)) if match.group(1) is not None else (2, match.group(2))
        end = headings[index + 1].start() if index + 1 < len(headings) else len(text)
        if level == 1 and re.match(r"(?:II|III|IV|V|VI|VII)\. ", heading):
            current_chapter = heading
            continue
        if level != 2 or heading in FRONT_HEADINGS or match.start() < front_end:
            continue
        source = parse_source(text[match.end():end])
        if not source:
            continue
        post_id, url, date, article_text = source
        articles.append({
            "chapter": current_chapter,
            "title": heading,
            "post_id": post_id,
            "article_id": f"mayaismagic-{post_id}",
            "url": url,
            "date": date,
            "text": article_text,
        })
    return articles


def source_links(article: dict[str, object]) -> list[dict[str, object]]:
    return list(article.get("source_links", [{key: article[key] for key in ("channel", "post_id", "url", "date")}]))


def unify_articles(primary: list[dict[str, object]], supplemental: list[dict[str, object]]) -> list[dict[str, object]]:
    """Keep a single reader-facing text while preserving every source URL."""
    canonical = {int(article["post_id"]): article for article in primary}
    for article in primary:
        prefix = str(article["chapter"]).split(".", 1)[0]
        article["chapter"] = next((chapter for chapter in CHAPTERS if chapter.startswith(prefix + ".")), str(article["chapter"]))
        article["channel"] = "Mayaismagic"
        article["source_links"] = [{key: article[key] for key in ("channel", "post_id", "url", "date")}]
    retained = list(primary)
    for article in supplemental:
        original = {key: article[key] for key in ("channel", "post_id", "url", "date")}
        canonical_id = SUPPLEMENTAL_CANONICAL.get(int(article["post_id"]))
        if canonical_id is not None:
            canonical[canonical_id]["source_links"].append(original)
            continue
        article["source_links"] = [original]
        retained.append(article)
    positions = {chapter: index for index, chapter in enumerate(CHAPTERS)}
    return sorted(retained, key=lambda article: (positions.get(str(article["chapter"]), 99), str(article["date"]), str(article["article_id"])))


def source_markdown(article: dict[str, object]) -> str:
    return " · ".join(
        f"{source['channel']}: пост [{source['post_id']}]({source['url']}); {source['date']}"
        for source in source_links(article)
    )


def write_unified_manuscript(articles: list[dict[str, object]], description: dict[str, object]) -> None:
    """Write the canonical reader manuscript; source archive files are never modified."""
    parts = ["# Maya Tradition", "", "## Описание традиции", "", f"*Источник: пост [{description['post_id']}]({description['url']}); {description['date']}.*", "", str(description["text"]), "", "# Содержание", ""]
    parts.extend(f"- {chapter}" for chapter in CHAPTERS)
    for chapter in CHAPTERS:
        parts.extend(["", f"# {chapter}"])
        for article in (item for item in articles if item["chapter"] == chapter):
            parts.extend(["", f"## {article['title']}", "", f"*Источники: {source_markdown(article)}.*", "", str(article["text"])])
    UNIFIED_MANUSCRIPT.write_text("\n".join(parts) + "\n", encoding="utf-8")


def normalized_text(text: str) -> str:
    return re.sub(r"[^\w]+", "", html.unescape(text).casefold())


def write_supporting_docs(articles: list[dict[str, object]]) -> None:
    """Regenerate reader-facing maps and the reproducible deduplication audit."""
    raw_sources = (HERE / "raw" / "posts.jsonl", SUPPLEMENTAL_INDEX)
    records = []
    for path in raw_sources:
        records.extend(json.loads(line) for line in path.read_text(encoding="utf-8").splitlines())
    exact: dict[str, list[dict[str, object]]] = defaultdict(list)
    for record in records:
        key = normalized_text(str(record["raw_text"]))
        if key:
            exact[key].append(record)
    groups = [group for group in exact.values() if len(group) > 1]
    report = ["# Deduplication report", "", "Scope: `raw/posts.jsonl` and `raw/templetherapy/TEMPLETHERAPY_MAYA_AZTEC_INDEX.jsonl`. Raw archives were read only.", "", f"- Source records audited: {len(records)}", f"- Normalized exact-duplicate groups: {len(groups)}", "- Cross-source exact relationships: 14", "- Near-duplicate decisions: 3", f"- Canonical reader articles retained: {len(articles)}", "", "## Exact duplicate groups", ""]
    for group in groups:
        links = ", ".join(f"[{item['channel']}:{item['post_id']}]({item['url']})" for item in group)
        report.append(f"- {links}")
    report.extend(["", "## Cross-source and near-duplicate decisions", "", "| Supplemental source | Canonical retained article | Decision and rationale |", "|---|---|---|"])
    for supplemental_id, canonical_id in sorted(SUPPLEMENTAL_CANONICAL.items()):
        kind = "exact copy after whitespace/entity normalization" if supplemental_id not in {2246, 2268, 2346} else {2246: "near copy; TempleTherapy adds a tail", 2268: "near copy; only a short closing sentence changes", 2346: "near copy; same material reordered"}[supplemental_id]
        report.append(f"| TempleTherapy:{supplemental_id} | Mayaismagic:{canonical_id} | {kind}; retain Mayaismagic text and link both sources. |")
    report.extend(["", "## Retention rule", "", "Each duplicate is represented once in the unified manuscript and reading editions. The canonical article keeps a link, date, and channel label for every duplicate source. Non-duplicate TempleTherapy entries are placed in the seven shared thematic chapters with their own labels, original URLs, and dates."])
    DEDUPLICATION_REPORT.write_text("\n".join(report) + "\n", encoding="utf-8")

    chapter_rows = ["# Chapter map", "", "The canonical manuscript is `manuscript/MAYA_TRADITION_UNIFIED.md`; it replaces the former TempleTherapy appendix in reader outputs.", ""]
    for chapter in CHAPTERS:
        members = [item for item in articles if item["chapter"] == chapter]
        chapter_rows.extend([f"## {chapter}", "", f"Articles: {len(members)}", ""])
        chapter_rows.extend(f"- `{item['article_id']}` — {item['title']}" for item in members)
        chapter_rows.append("")
    (HERE / "manuscript" / "CHAPTER_MAP.md").write_text("\n".join(chapter_rows), encoding="utf-8")
    (HERE / "CONTENT_MAP.md").write_text("\n".join(["# Content map", "", "All reader articles are listed in the chapter map; IDs are namespace-safe (`mayaismagic-<id>` or `templetherapy-<id>`).", "", f"Canonical article count: {len(articles)}."]) + "\n", encoding="utf-8")
    (HERE / "manuscript" / "COVERAGE.md").write_text(f"# Coverage\n\n- Primary curated articles: 81\n- Supplemental source posts: 29\n- Supplemental duplicates consolidated: 17\n- Final canonical reader articles: {len(articles)}\n", encoding="utf-8")
    (HERE / "manuscript" / "SOURCE_NOTES.md").write_text("# Source notes\n\nMayaismagic is the primary Telegram export. TempleTherapy is a separately-labelled supplemental public source. Duplicate source URLs are preserved on the canonical retained article; no archive content was altered.\n", encoding="utf-8")
    (HERE / "FACT_CHECK.md").write_text("# Fact-check boundary\n\nThis is a source-backed reading edition, not an independent historical fact-check. Editorial work is limited to placement, deduplication, and source attribution; claims in posts remain attributed to their original channel.\n", encoding="utf-8")


def article_html(article: dict[str, object], primary: str | None) -> str:
    photo = ""
    if primary:
        photo = (f'<figure class="post-media"><img class="post-photo-main" src="../{html.escape(primary)}" '
                 f'alt="Источник: пост {article["post_id"]}" loading="lazy"></figure>')
    text = "<br>\\n".join(html.escape(line.rstrip()) for line in str(article["text"]).splitlines())
    sources = "".join(f'<a href="{html.escape(str(source["url"]))}">{html.escape(str(source["channel"]))} · пост {source["post_id"]}</a>' for source in source_links(article))
    return f'''<article class="post" id="{html.escape(str(article["article_id"]))}">
  <div class="chapter-token">{html.escape(str(article["chapter"]))}</div>
  <h2>{html.escape(str(article["title"]))}</h2>
  <div class="meta"><span>Источники:</span>{sources}<span>Дата: {html.escape(str(article["date"]))}</span></div>
{photo}
  <div class="text">{text}</div>
</article>'''


def chapter_id(chapter: str) -> str:
    return f"chapter-{chapter.split('.', 1)[0].lower()}"


def meta_html(item: dict[str, object]) -> str:
    return (f'<div class="meta"><span>Источник: {html.escape(str(item.get("channel", "mayaismagic")))} · пост {item["post_id"]}</span><span>Дата: {html.escape(str(item["date"]))}</span>'
            f'<a href="{item["url"]}">{item["url"]}</a></div>')


def build_html(articles: list[dict[str, object]], media: dict[int, list[str]], description: dict[str, object]) -> None:
    sections: list[str] = []
    chapters: list[str] = []
    chapter = None
    for article in articles:
        if article["chapter"] != chapter:
            chapter = str(article["chapter"])
            chapters.append(chapter)
            sections.append(f'<h1 class="chapter" id="{chapter_id(chapter)}">{html.escape(chapter)}</h1>')
        if article.get("channel") == "TempleTherapy":
            primary = next((f"media/templetherapy/{path.name}" for path in sorted(SUPPLEMENTAL_MEDIA.glob(f"post-{article['post_id']}-*"))), None)
        else:
            source_media = (media.get(int(article["post_id"])) or [None])[0]
            primary = f"raw/{source_media}" if source_media else None
        sections.append(article_html(article, primary))
    toc = "".join(f'<li><a href="#{chapter_id(chapter)}">{html.escape(chapter)}</a></li>' for chapter in chapters)
    description_text = "<br>\n".join(html.escape(line.rstrip()) for line in str(description["text"]).splitlines())
    document = "\n".join(sections)
    HTML_OUT.write_text(f'''<!doctype html>
<html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>Maya Tradition: методология источникового чтения</title>
<style>
:root{{--ink:#{INK};--wine:#{WARM};--gold:#{GOLD};--paper:#{CREAM};--line:#ddc9b6}} *{{box-sizing:border-box}}
body{{margin:0;background:#efe5d8;color:var(--ink);font:19px/1.75 Georgia,"Times New Roman",serif}} main{{max-width:980px;margin:auto;padding:36px 20px 80px}}
.cover{{background:linear-gradient(135deg,#4d2117,var(--wine));color:#fff7ec;border-radius:18px;padding:42px 38px;margin-bottom:28px}} .eyebrow,.chapter-token{{font:700 12px/1.2 Arial,sans-serif;letter-spacing:.12em;text-transform:uppercase;color:var(--gold)}} .cover .eyebrow{{color:#f2c979}} .cover h1{{font-size:clamp(34px,5vw,54px);line-height:1.08;margin:.3em 0}} .cover p{{max-width:720px;line-height:1.75;margin:0}}
.front-card,.toc{{background:#fffdf9;border:1px solid var(--line);border-radius:14px;padding:24px 26px;margin:18px 0}} .front-card h2,.toc h2{{font-size:28px;line-height:1.25;margin:0 0 10px;color:var(--wine)}} .toc ol{{list-style:none;padding:0;margin:0;display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:10px}} .toc a{{display:block;min-height:48px;padding:11px 14px;border-radius:9px;background:#f8eee2;color:var(--wine);font:700 17px/1.35 Arial,sans-serif;text-decoration:none}}
.chapter{{scroll-margin-top:16px;font-size:29px;line-height:1.25;color:var(--wine);margin:48px 0 16px;padding-bottom:9px;border-bottom:2px solid var(--gold)}} .post{{background:#fffdf9;border:1px solid var(--line);border-radius:14px;padding:23px 25px;margin:18px 0;display:flow-root;break-before:page;page-break-before:always}} .post h2{{font-size:27px;line-height:1.27;margin:7px 0 11px;color:var(--wine)}} .meta{{display:flex;gap:8px 13px;flex-wrap:wrap;font:14px/1.5 Arial,sans-serif;color:#{MUTED};padding:10px 0 14px;border-top:1px solid #eadacc;border-bottom:1px solid #eadacc;margin-bottom:16px}} .meta a{{color:var(--wine);overflow-wrap:anywhere}} .post-media{{float:right;width:min(36%,310px);margin:0 0 15px 24px}} .post-photo-main{{display:block;width:100%;height:auto;border-radius:10px;border:1px solid #d4b89e}} .text{{white-space:normal;font-size:1rem;line-height:1.75}}
@media(max-width:700px){{body{{font-size:18px;line-height:1.75}}main{{padding:18px 12px 48px}}.cover{{padding:29px 22px}}.front-card,.toc,.post{{padding:19px 17px}}.toc ol{{grid-template-columns:1fr;gap:9px}}.toc a{{min-height:52px;font-size:17px;padding:13px 14px}}.post-media{{float:none;width:100%;margin:0 0 15px}}.post h2{{font-size:25px}}}} @media print{{body{{background:#fff}}main{{max-width:none;padding:0}}.cover{{border-radius:0;break-after:page}}.toc{{break-after:page}}.post{{border-radius:0;margin:0;min-height:92vh}}}}
</style></head><body><main><header class="cover"><div class="eyebrow">Reading edition · local Telegram export</div><h1>Maya Tradition</h1><p>Методология источникового чтения. Редакционная компоновка сохранённых текстов без фактологического дополнения.</p></header><section class="front-card"><h2>Описание традиции</h2>{meta_html(description)}<div class="text">{description_text}</div></section><nav class="toc" aria-label="Содержание"><h2>Содержание</h2><ol>{toc}</ol></nav>{document}</main></body></html>''', encoding="utf-8")

def shade(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), value); tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); margins = tcPr.first_child_found_in("w:tcMar")
    if margins is None: margins = OxmlElement("w:tcMar"); tcPr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); margins.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def add_run(paragraph, text, size=13, bold=False, color=INK):
    run = paragraph.add_run(text); run.bold = bold; run.font.name = "Georgia"; run._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia"); run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); return run


def set_paragraph(paragraph, before=0, after=6, line_spacing=1.35):
    fmt = paragraph.paragraph_format; fmt.space_before = Pt(before); fmt.space_after = Pt(after); fmt.line_spacing = line_spacing


def build_docx(articles: list[dict[str, object]], media: dict[int, list[str]], description: dict[str, object]) -> None:
    doc = Document(); sec = doc.sections[0]; sec.top_margin = Inches(.68); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.72); sec.right_margin = Inches(.72)
    styles = doc.styles; styles["Normal"].font.name = "Georgia"; styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia"); styles["Normal"].font.size = Pt(13); styles["Normal"].paragraph_format.line_spacing = 1.35
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(title, after=7); add_run(title, "MAYA TRADITION", 28, True, WARM)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(subtitle, after=7); add_run(subtitle, "Методология источникового чтения", 15, True, GOLD)
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_paragraph(note, after=15); add_run(note, "Локальный Telegram-экспорт · источник-ориентированная редакционная компоновка", 10.5, False, MUTED)
    heading = doc.add_paragraph(); set_paragraph(heading, after=6); add_run(heading, "Описание традиции", 18, True, WARM)
    meta = doc.add_paragraph(); set_paragraph(meta, after=6); add_run(meta, f"Источник: пост {description['post_id']} · {description['date']}\n", 10.5, True, GOLD); add_run(meta, str(description["url"]), 10.5, False, WARM)
    desc = doc.add_paragraph(); set_paragraph(desc, after=14); add_run(desc, str(description["text"]), 13)
    toc_heading = doc.add_paragraph(); set_paragraph(toc_heading, after=6); add_run(toc_heading, "Содержание", 18, True, WARM)
    seen: list[str] = []
    for article in articles:
        chapter_name = str(article["chapter"])
        if chapter_name not in seen:
            seen.append(chapter_name)
            toc_line = doc.add_paragraph(); set_paragraph(toc_line, after=4); add_run(toc_line, chapter_name, 13, True, WARM)
    chapter = None
    for article in articles:
        if article["chapter"] != chapter:
            chapter = article["chapter"]
        doc.add_page_break()
        token = doc.add_table(rows=1, cols=1); token.autofit = False; cell = token.cell(0,0); shade(cell, CREAM); set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; set_paragraph(p, after=0); add_run(p, str(chapter).upper(), 10, True, GOLD)
        table = doc.add_table(rows=1, cols=2); table.autofit = False; table.columns[0].width = Inches(4.3); table.columns[1].width = Inches(2.0)
        left, right = table.rows[0].cells; set_cell_margins(left); set_cell_margins(right); right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = left.paragraphs[0]; set_paragraph(p, after=7); add_run(p, str(article["title"]), 18, True, WARM)
        meta = left.add_paragraph(); set_paragraph(meta, after=0)
        for source in source_links(article):
            add_run(meta, f"Источник: {source['channel']} · пост {source['post_id']}\n", 10.5, True, GOLD)
            add_run(meta, f"Дата: {source['date']}\n", 10.5, False, MUTED)
            add_run(meta, str(source["url"]) + "\n", 10.5, False, WARM)
        if article.get("channel") == "TempleTherapy":
            primary_path = next(iter(sorted(SUPPLEMENTAL_MEDIA.glob(f"post-{article['post_id']}-*"))), None)
        else:
            primary = (media.get(int(article["post_id"])) or [None])[0]
            primary_path = HERE / "raw" / primary if primary else None
        if primary_path and primary_path.exists():
            p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.add_run().add_picture(str(primary_path), width=Inches(1.85))
        body = doc.add_paragraph(); set_paragraph(body, before=10, after=6); add_run(body, str(article["text"]), 13)
    doc.save(DOCX_OUT)


def main() -> None:
    if Document is None:
        raise SystemExit("Building the DOCX edition requires python-docx")
    OUT.mkdir(exist_ok=True)
    text = MANUSCRIPT.read_text(encoding="utf-8")
    primary_articles, supplementary_articles = parse_articles(), parse_supplemental_articles()
    articles, media, description = unify_articles(primary_articles, supplementary_articles), parse_media(), parse_front_description(text)
    if len(primary_articles) != 81 or len(supplementary_articles) != 29:
        raise SystemExit(f"Expected 81 primary and 29 supplemental reading articles, found {len(primary_articles)} and {len(supplementary_articles)}")
    write_unified_manuscript(articles, description)
    write_supporting_docs(articles)
    build_html(articles, media, description)
    build_docx(articles, media, description)
    print(f"wrote unified manuscript, maps, {HTML_OUT} and {DOCX_OUT} ({len(articles)} canonical articles)")


if __name__ == "__main__":
    main()
